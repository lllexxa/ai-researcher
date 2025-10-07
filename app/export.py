from __future__ import annotations

import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document

from .cite import Reference


@dataclass
class ExportResult:
    markdown_path: Path
    docx_path: Path


def _generate_filename(prefix: str, suffix: str) -> str:
    return f"{prefix}-{uuid.uuid4().hex}{suffix}"


def _write_markdown(content: str, target_dir: Path) -> Path:
    filename = _generate_filename("review", ".md")
    path = target_dir / filename
    path.write_text(content, encoding="utf-8")
    return path


def _line_as_docx(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        document.add_paragraph("")
        return
    if stripped.startswith("#"):
        level = stripped.count("#")
        text = stripped[level:].strip()
        level = max(1, min(level, 5))
        document.add_heading(text, level=level)
        return
    if stripped.startswith(('- ', '* ')):
        document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return
    if stripped[0].isdigit() and stripped[1:3] == ". ":  # simple ordered list
        document.add_paragraph(stripped[3:].strip(), style="List Number")
        return
    document.add_paragraph(stripped)


def _write_docx(content: str, references: Iterable[Reference], target_dir: Path) -> Path:
    filename = _generate_filename("review", ".docx")
    path = target_dir / filename
    document = Document()
    for line in content.splitlines():
        _line_as_docx(document, line)
    if references:
        document.add_heading("References", level=1)
        for ref in references:
            document.add_paragraph(f"[{ref.index}] {ref.text}")
    document.save(path)
    return path


def export_review(markdown: str, references: Iterable[Reference], downloads_dir: Path) -> ExportResult:
    downloads_dir.mkdir(parents=True, exist_ok=True)
    md_path = _write_markdown(markdown, downloads_dir)
    docx_path = _write_docx(markdown, references, downloads_dir)
    return ExportResult(markdown_path=md_path, docx_path=docx_path)
